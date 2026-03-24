from typing import List, Dict, Tuple

try:
    import docx
except ImportError:
    docx = None

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

def extract_text_with_pages(file_path: str) -> List[Dict[str, str]]:
    """
    通过 pdfplumber 或 python-docx 提取文档内容并带有原始页码标记
    返回结构: [{"page_num": "1", "text": "第一页原文内容"}, ...]
    """
    ext = file_path.lower().split('.')[-1]
    if ext == 'pdf':
        return _parse_pdf(file_path)
    elif ext in ['doc', 'docx']:
        return _parse_docx(file_path)
    else:
        raise ValueError(f"暂不支持解析的文件格式后缀: {ext}")

def extract_native_toc(file_path: str) -> Tuple[str, int]:
    """
    高速第一层漏斗：尝试从文档的原生元数据（如 PDF 的 Outlines 书签，Word 的 Heading 样式）中直接秒取目录树。
    返回: (Markdown 格式的目录树，正式正文在此次文档切片组中的起始索引)
    """
    ext = file_path.lower().split('.')[-1]
    
    if ext == 'pdf':
        if not fitz:
            return "", 0 # 没有安装 PyMuPDF 则放弃原生提取，降级由 LLM 处理
        try:
            doc = fitz.open(file_path)
            toc = doc.get_toc() # 返回格式: [[level, title, page_num], ...]
            if not toc:
                return "", 0
            
            # 使用提取到的最低页码书签定位作为文档内容开始的地方！
            # page_num 是基于 1 的物理页，转换成 list index 要 -1
            min_page = min([item[2] for item in toc]) if toc else 1
            start_idx = max(0, int(min_page) - 1)
            
            toc_md_lines = []
            for item in toc:
                level = item[0]
                title = item[1].strip()
                page_num = item[2]
                prefix = "#" * level
                toc_md_lines.append(f"{prefix} {title}\n页码 {page_num}")
            return "\n".join(toc_md_lines), start_idx
        except Exception:
            return "", 0
            
    elif ext in ['doc', 'docx']:
        if not docx:
            return "", 0
        try:
            doc = docx.Document(file_path)
            toc_lines = []
            first_heading_idx = 0
            found_first = False
            for i, p in enumerate(doc.paragraphs):
                if p.style and p.style.name and p.style.name.startswith('Heading'):
                    if not found_first:
                        first_heading_idx = i // 20
                        found_first = True
                    level_str = p.style.name.replace('Heading', '').strip()
                    level = int(level_str) if level_str.isdigit() else 1
                    prefix = "#" * level
                    # 根据前文的分块切开规则，20 段约为一页，利用虚拟分页索引赋予其页码
                    virtual_page = (i // 20) + 1
                    toc_lines.append(f"{prefix} {p.text.strip()}\n页码 {virtual_page}")
            return "\n".join(toc_lines), first_heading_idx
        except Exception:
            return "", 0
            
    return "", 0

def _parse_pdf(file_path: str) -> List[Dict[str, str]]:
    pages_data = []
    with pdfplumber.open(file_path) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text()
            if text and text.strip():
                pages_data.append({
                    "page_num": str(i + 1),
                    "text": text.strip()
                })
    return pages_data

def _parse_docx(file_path: str) -> List[Dict[str, str]]:
    if not docx:
        raise ImportError("请安装 python-docx 库以解析 Word 文档")
    doc = docx.Document(file_path)
    
    # 简单实现：将 Docx 合并处理，因无严格物理页码，可按一定段落数量分块挂接虚拟页码
    # 此处暂定以每 20 个段落视为 1 个"虚拟页码"片段
    pages_data = []
    current_page = 1
    current_text = []
    
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            current_text.append(p.text.strip())
        if (i > 0 and i % 20 == 0) or i == len(doc.paragraphs) - 1:
            if current_text:
                pages_data.append({
                    "page_num": str(current_page),
                    "text": "\n".join(current_text)
                })
                current_page += 1
                current_text = []
                
    return pages_data
